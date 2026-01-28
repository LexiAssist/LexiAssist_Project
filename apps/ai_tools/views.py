import logging
import os
from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.conf import settings
from apps.classroom.models import UserActivity

import io
from gtts import gTTS
from django.http import FileResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response

# Consolidated imports from services
from .models import ChatSession, ChatMessage
from .services import (
    get_ai_response,
    generate_flashcards,
    generate_quiz_data,
    extract_text_from_pdf
)

from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

# --- AI ENGINE IMPORTS (For RAG) ---
from ai_engine.doc_loader import load_and_split_pdf
from ai_engine.embeddings import create_vector_store
from ai_engine.rag import workflow
import ai_engine.retriever as retriever_module

logger = logging.getLogger(__name__)


# ============================
# HELPER FUNCTION: Extract Text from Any Document
# ============================
def extract_text_from_document(uploaded_file):
    """
    Extract text from PDF, Word, or TXT files.
    Returns (success, text_or_error_message)
    """
    filename_lower = uploaded_file.name.lower()

    try:
        if filename_lower.endswith('.pdf'):
            text = extract_text_from_pdf(uploaded_file)
            print(f"✅ Extracted from PDF: {len(text)} chars")
            return True, text

        elif filename_lower.endswith('.txt'):
            text = uploaded_file.read().decode('utf-8')
            print(f"✅ Extracted from TXT: {len(text)} chars")
            return True, text

        elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
            try:
                from docx import Document
                doc = Document(uploaded_file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                print(f"✅ Extracted from Word: {len(text)} chars")
                return True, text
            except ImportError:
                return False, "Word document support not installed. Please contact administrator."
            except Exception as e:
                return False, f"Error reading Word document: {str(e)}"
        else:
            return False, "Unsupported file type. Please upload PDF, Word (DOC/DOCX), or TXT."

    except Exception as e:
        return False, f"Error extracting text: {str(e)}"


# ============================
# 1. CHAT ASSISTANT
# ============================
@login_required
def chat_assistant_view(request, session_id=None):
    # 1. Get all previous chats for the sidebar (INCLUDING Study Buddy sessions)
    my_sessions = ChatSession.objects.filter(user=request.user).order_by('-created_at')

    # 2. If a specific session is requested (clicking history), load its messages
    current_session = None
    existing_messages = []

    if session_id:
        current_session = get_object_or_404(ChatSession, id=session_id, user=request.user)
        existing_messages = current_session.messages.all().order_by('created_at')

    # 3. Handle the AJAX Chat (When you click Send)
    if request.method == "POST":
        user_query = request.POST.get('user_query', '').strip()
        uploaded_image = request.FILES.get('image')
        active_session_id = request.POST.get('session_id')

        if not user_query and not uploaded_image:
            return JsonResponse({'status': 'error', 'reply': 'Empty message.'})

        try:
            # A. Get or Create the Session
            if active_session_id:
                session = ChatSession.objects.get(id=active_session_id, user=request.user)
            else:
                title_text = user_query[:30] + "..." if user_query else "Image Analysis"
                session = ChatSession.objects.create(user=request.user, title=title_text)

            # B. Save USER Message
            ChatMessage.objects.create(session=session, is_user=True, text=user_query)

            # C. Get AI Response
            ai_reply = get_ai_response(text=user_query, image=uploaded_image)

            # D. Save AI Message
            ChatMessage.objects.create(session=session, is_user=False, text=ai_reply)

            return JsonResponse({
                'status': 'success',
                'reply': ai_reply,
                'session_id': session.id,
                'session_title': session.title
            })

        except Exception as e:
            logger.error(f"Error: {e}")
            return JsonResponse({'status': 'error', 'reply': "An error occurred."})

    context = {
        'chat_sessions': my_sessions,
        'active_session': current_session,
        'existing_messages': existing_messages
    }
    return render(request, 'chat_assistant.html', context)


# ============================
# 2. FLASHCARD GENERATOR
# ============================
def flashcard_generator_view(request):
    if request.method == "POST":
        topic = request.POST.get('topic')
        num_cards = request.POST.get('num_cards', 5)
        upload_file = request.FILES.get('file')

        # 1. Get Text from File or Paste
        content_text = ""
        if upload_file:
            success, result = extract_text_from_document(upload_file)
            if success:
                content_text = result
            else:
                content_text = f"Error: {result}"

        # 2. Call Gemini
        cards_data = generate_flashcards(topic, content_text, num_cards)

        return JsonResponse({'status': 'success', 'cards': cards_data})

    return render(request, 'flashcards.html')


# ============================
# 3. QUIZ GENERATOR
# ============================
def quizzes_view(request):
    if request.method == "POST":
        topic = request.POST.get('topic')
        num_questions = 5  # Default to 5

        # 1. Find the user's most recent uploaded PDF/Doc
        last_activity = UserActivity.objects.filter(
            user=request.user
        ).exclude(file_name='').order_by('-timestamp').first()

        context_text = ""
        if last_activity:
            file_path = os.path.join(settings.MEDIA_ROOT, last_activity.file_name)
            # Check if file actually exists on disk
            if os.path.exists(file_path):
                filename_lower = file_path.lower()
                if filename_lower.endswith('.pdf'):
                    with open(file_path, 'rb') as f:
                        context_text = extract_text_from_pdf(f)
                elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        context_text = ""
                        for paragraph in doc.paragraphs:
                            context_text += paragraph.text + "\n"
                    except:
                        pass

        # If no file found, use the topic as context
        if not context_text:
            context_text = f"General knowledge about {topic}"

        # 2. Call AI
        quiz_data = generate_quiz_data(topic, context_text, num_questions)

        return JsonResponse({'status': 'success', 'quiz': quiz_data})

    return render(request, 'quizzes.html')


# ============================
# 4. RAG / STUDY BUDDY API
# ============================

@csrf_exempt
def upload_note_api(request):
    """
    API View to handle PDF uploads and process them for the AI.
    """
    if request.method == 'POST' and request.FILES.get('document'):
        uploaded_file = request.FILES['document']

        # 1. Save the file temporarily
        file_path = default_storage.save(f"temp_notes/{uploaded_file.name}", ContentFile(uploaded_file.read()))
        full_path = os.path.join(default_storage.location, file_path)

        try:
            # 2. Load and Split the PDF
            print(f"\n{'='*60}")
            print(f"📄 UPLOAD NOTE API - Processing file: {full_path}")
            print(f"{'='*60}\n")

            documents = load_and_split_pdf(full_path)
            print(f"✅ PDF loaded and split into {len(documents)} chunks")

            # 3. Create Vector Store (The Brain)
            print("🧠 Creating embeddings... (this might take a moment)")
            vector_store = create_vector_store(documents)
            print("✅ Vector store created and saved to disk")

            # 4. Update the global variable in retriever.py so the AI uses THIS document
            retriever_module.active_vector_store = vector_store
            print("✅ Active vector store updated\n")

            # ===== CREATE CHAT SESSION FOR THIS DOCUMENT =====
            if request.user.is_authenticated:
                # Extract clean document name
                doc_name = uploaded_file.name.replace('.pdf', '').replace('_', ' ')[:40]
                session_title = f"📚 Study: {doc_name}"

                # Create NEW session for this document
                study_session = ChatSession.objects.create(
                    user=request.user,
                    title=session_title
                )

                print(f"✅ Created new Study Buddy session: {session_title}")

                # Store session ID so chat_with_ai_api can use it
                request.session['current_study_session_id'] = study_session.id

            return JsonResponse({
                "status": "success",
                "message": "Note processed successfully! You can now chat."
            })

        except Exception as e:
            logger.error(f"Error processing RAG file: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

    return JsonResponse({"status": "error", "message": "No file uploaded"}, status=400)


@csrf_exempt
def chat_with_ai_api(request):
    """
    API View to handle RAG chat messages (Specific to the uploaded document).
    """
    if request.method == 'POST':
        user_message = request.POST.get('message')

        if not user_message:
            return JsonResponse({"error": "Message is required"}, status=400)

        print(f"\n{'='*60}")
        print(f"💬 CHAT WITH AI API")
        print(f"📩 User Message: {user_message}")
        print(f"{'='*60}\n")

        # Check if FAISS index exists
        faiss_path = os.path.join(settings.BASE_DIR, "faiss_index")
        print(f"📁 Checking for FAISS index at: {faiss_path}")
        print(f"📁 Index exists: {os.path.exists(faiss_path)}")

        if not os.path.exists(faiss_path):
            print("❌ No FAISS index found")
            return JsonResponse({
                "answer": "No document uploaded yet. Please upload a PDF file first."
            })

        try:
            # Import the retrieve_data tool directly
            from ai_engine.retriever import retrieve_data

            print("🔄 Calling retrieve_data directly...")

            # Call the retriever directly
            ai_answer = retrieve_data.invoke({"query": user_message})

            print(f"\n✅ Got answer from retrieve_data")
            print(f"📝 Answer preview: {ai_answer[:200]}...\n")

            # ===== SAVE TO THE CURRENT STUDY SESSION =====
            if request.user.is_authenticated:
                # Get the session ID that was created when document was uploaded
                study_session_id = request.session.get('current_study_session_id')

                if study_session_id:
                    try:
                        study_session = ChatSession.objects.get(
                            id=study_session_id,
                            user=request.user
                        )
                        print(f"✅ Using session: {study_session.title}")
                    except ChatSession.DoesNotExist:
                        # Fallback: create new session if not found
                        study_session = ChatSession.objects.create(
                            user=request.user,
                            title="📚 Study Session"
                        )
                        request.session['current_study_session_id'] = study_session.id
                        print(f"⚠️ Created fallback session")
                else:
                    # No session in memory, create new one
                    study_session = ChatSession.objects.create(
                        user=request.user,
                        title="📚 Study Session"
                    )
                    request.session['current_study_session_id'] = study_session.id
                    print(f"⚠️ No session found, created new one")

                # Save user message
                ChatMessage.objects.create(
                    session=study_session,
                    is_user=True,
                    text=user_message
                )

                # Save AI response
                ChatMessage.objects.create(
                    session=study_session,
                    is_user=False,
                    text=ai_answer
                )

                print(f"✅ Messages saved to session: {study_session.title}")

            return JsonResponse({"answer": ai_answer})

        except Exception as e:
            logger.error(f"RAG Chat Error: {e}")
            import traceback
            traceback.print_exc()

            return JsonResponse({
                "answer": f"I'm having trouble reading that document. Error: {str(e)}"
            })

    return JsonResponse({"error": "Invalid request"}, status=400)


# ============================
# 5. TEXT TO SPEECH DASHBOARD
# ============================
@login_required
def tts_dashboard_view(request):
    """
    Handle file upload and extract text for Text-to-Speech.
    """
    print(f"\n{'='*60}")
    print(f"📋 TTS DASHBOARD VIEW CALLED")
    print(f"Method: {request.method}")
    print(f"Files: {request.FILES}")
    print(f"{'='*60}\n")

    if request.method == 'POST' and request.FILES.get('document'):
        uploaded_file = request.FILES['document']

        try:
            # Extract text using helper function
            success, result = extract_text_from_document(uploaded_file)

            if not success:
                return JsonResponse({
                    'status': 'error',
                    'message': result
                }, status=400)

            extracted_text = result

            # Split text into sentences for highlighting
            import re
            sentences = re.split(r'(?<=[.!?])\s+', extracted_text)
            sentences = [s.strip() for s in sentences if s.strip()]

            # Pass data to the reader page
            context = {
                'sentences': sentences,
                'filename': uploaded_file.name,
                'full_text': extracted_text
            }

            return render(request, 'text2speech_reader.html', context)

        except Exception as e:
            logger.error(f"TTS Upload Error: {e}")
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            }, status=500)

    # GET request - show upload page
    return render(request, 'text2speech.html')


# ============================
# 6. TEXT TO SPEECH API
# ============================
@api_view(['POST'])
def text_to_speech(request):
    """
    Converts text to speech and streams the audio file back.
    """
    text = request.data.get('text')
    language = request.data.get('language', 'en')

    if not text:
        return Response({"error": "Text input cannot be empty"}, status=400)

    try:
        # Create a memory buffer
        mp3_fp = io.BytesIO()

        # Generate speech using gTTS
        tts = gTTS(text=text, lang=language, slow=False)
        tts.write_to_fp(mp3_fp)

        # Reset buffer pointer
        mp3_fp.seek(0)

        # Return as streaming file response
        response = FileResponse(mp3_fp, content_type='audio/mpeg')
        response['Content-Disposition'] = 'attachment; filename="speech.mp3"'
        return response

    except Exception as e:
        return Response({"error": str(e)}, status=500)


# ============================
# 7. READING ASSISTANT (AI TEXT SIMPLIFICATION)
# ============================
@login_required
def reading_assistant_view(request):
    """
    Handle file upload and simplify text for dyslexic readers.
    """
    if request.method == 'POST' and request.FILES.get('document'):
        uploaded_file = request.FILES['document']

        print(f"\n{'='*60}")
        print(f"📖 READING ASSISTANT - Processing: {uploaded_file.name}")
        print(f"{'='*60}\n")

        try:
            # Extract text using helper function
            success, result = extract_text_from_document(uploaded_file)

            if not success:
                return JsonResponse({
                    'status': 'error',
                    'message': result
                }, status=400)

            original_text = result

            if len(original_text.strip()) < 50:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Document appears to be empty or too short.'
                }, status=400)

            print(f"✅ Original text extracted ({len(original_text)} chars)")

            # Use Gemini to simplify the text
            import google.generativeai as genai
            genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

            simplification_prompt = f"""You are a reading assistant for students with dyslexia.

Your task: Simplify this text to make it easier to read.

Rules:
1. Use shorter sentences (maximum 15 words per sentence)
2. Replace complex words with simpler alternatives
3. Break paragraphs into smaller chunks
4. Keep the same meaning and important information
5. Use clear, direct language

Original Text:
{original_text[:3000]}

Provide the simplified version:"""

            print("🤖 Calling Gemini to simplify text...")
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(simplification_prompt)
            simplified_text = response.text

            print(f"✅ Text simplified ({len(simplified_text)} chars)")

            # Extract vocabulary words
            vocab_prompt = f"""From this text, identify 5-7 difficult words that a dyslexic student might struggle with.

For each word, provide:
1. The word
2. A simple, one-sentence definition (10 words or less)

Text:
{original_text[:2000]}

Format your response as a JSON array:
[
  {{"word": "scrutinized", "definition": "looked at very closely"}},
  {{"word": "ideology", "definition": "a set of beliefs or ideas"}}
]

Return ONLY the JSON array, no other text."""

            print("📚 Extracting vocabulary...")
            vocab_response = model.generate_content(vocab_prompt)
            vocab_text = vocab_response.text.strip()

            # Clean and parse vocab JSON
            vocab_text = vocab_text.replace('```json', '').replace('```', '').strip()

            import json
            try:
                vocab_list = json.loads(vocab_text)
                print(f"✅ Found {len(vocab_list)} vocabulary words")
            except:
                vocab_list = []
                print("⚠️ Could not parse vocab, using empty list")

            # Pass everything to the reader page
            context = {
                'filename': uploaded_file.name,
                'original_text': original_text,
                'simplified_text': simplified_text,
                'vocab_list': vocab_list
            }

            return render(request, 'reading_reader.html', context)

        except Exception as e:
            logger.error(f"Reading Assistant Error: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            }, status=500)

    # GET request - show upload page
    return render(request, 'readingass.html')
